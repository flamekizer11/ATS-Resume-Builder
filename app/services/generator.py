"""
Resume Generator Service
Creates ATS-friendly DOCX resumes from structured data
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# This module takes structured resume data (as a dictionary) and generates a well-formatted DOCX resume.
def generate_ats_resume(data: dict, filename: str):
    """
    Generate ATS-friendly DOCX resume
    
    Args:
        data: Structured resume data
        filename: Output file path
    """
    doc = Document()
    
    # Set narrow margins which is ATS Friendly
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # HEADER - Name and Contact
    personal = data.get('personal', {})
    name = doc.add_paragraph()
    name_run = name.add_run(personal.get('name', 'Your Name'))
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = doc.add_paragraph()
    contact_parts = []
    
    if personal.get('email'):
        contact_parts.append(personal['email'])
    if personal.get('phone'):
        contact_parts.append(personal['phone'])
    if personal.get('linkedin'):
        contact_parts.append(personal['linkedin'])
    if personal.get('location'):
        contact_parts.append(personal['location'])
    
    contact_text = ' | '.join(contact_parts)
    contact.add_run(contact_text)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.runs[0].font.size = Pt(10)
    
    # For spacing between header and next section
    doc.add_paragraph()
    
    # PROFESSIONAL SUMMARY, TECHNICAL SKILLS and EXPERIENCE
    if data.get('summary'):
        add_section(doc, 'PROFESSIONAL SUMMARY', data['summary'])
    
    if data.get('skills'):
        skills_text = ' • '.join(data['skills'])
        add_section(doc, 'TECHNICAL SKILLS', skills_text)

    if data.get('experience'):
        heading = doc.add_paragraph()
        heading_run = heading.add_run('EXPERIENCE')
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        
        for exp in data['experience']:
            # Job Title and Company
            exp_header = doc.add_paragraph()
            title_company = f"{exp.get('title', '')} | {exp.get('company', '')}"
            title_run = exp_header.add_run(title_company)
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            # Duration
            if exp.get('duration'):
                duration_para = doc.add_paragraph()
                duration_run = duration_para.add_run(exp.get('duration', ''))
                duration_run.italic = True
                duration_run.font.size = Pt(10)
            
            # Bullet points
            for point in exp.get('points', []):
                p = doc.add_paragraph(point, style='List Bullet')
                p.runs[0].font.size = Pt(10)
            doc.add_paragraph()
    
    # PROJECTS 
    if data.get('projects'):
        heading = doc.add_paragraph()
        heading_run = heading.add_run('PROJECTS')
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        
        for proj in data['projects']:
            # Project Name
            proj_header = doc.add_paragraph()
            proj_name_run = proj_header.add_run(proj.get('name', ''))
            proj_name_run.bold = True
            proj_name_run.font.size = Pt(11)
            
            # Technologies
            if proj.get('tech'):
                tech_para = doc.add_paragraph()
                tech_text = f"Technologies: {', '.join(proj['tech'])}"
                tech_run = tech_para.add_run(tech_text)
                tech_run.italic = True
                tech_run.font.size = Pt(10)
            
            # Description
            if proj.get('description'):
                desc_para = doc.add_paragraph(proj['description'])
                desc_para.runs[0].font.size = Pt(10)
            
            # Bullet points
            for point in proj.get('points', []):
                p = doc.add_paragraph(point, style='List Bullet')
                p.runs[0].font.size = Pt(10)
            
            # Add spacing between projects
            doc.add_paragraph()
    
    # EDUCATION
    if data.get('education'):
        heading = doc.add_paragraph()
        heading_run = heading.add_run('EDUCATION')
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        
        for edu in data['education']:
            # Degree and Institution
            edu_p = doc.add_paragraph()
            degree_inst = f"{edu.get('degree', '')} | {edu.get('institution', '')}"
            degree_run = edu_p.add_run(degree_inst)
            degree_run.bold = True
            degree_run.font.size = Pt(11)
            
            # Year and GPA
            details = []
            if edu.get('year'):
                details.append(edu['year'])
            if edu.get('gpa'):
                details.append(f"GPA: {edu['gpa']}")
            
            if details:
                detail_para = doc.add_paragraph(' | '.join(details))
                detail_para.runs[0].font.size = Pt(10)
            
            # Add spacing
            doc.add_paragraph()
    
    # CERTIFICATIONS and ACHIEVEMENTS
    if data.get('certifications'):
        cert_text = ' • '.join(data['certifications'])
        add_section(doc, 'CERTIFICATIONS', cert_text)

    if data.get('achievements'):
        heading = doc.add_paragraph()
        heading_run = heading.add_run('ACHIEVEMENTS')
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        
        for achievement in data['achievements']:
            p = doc.add_paragraph(achievement, style='List Bullet')
            p.runs[0].font.size = Pt(10)
    
    # Save document
    doc.save(filename)


def add_section(doc, title: str, content: str):
    """
    Add a section with heading and content
    
    Args:
        doc: Document object
        title: Section title
        content: Section content
    """
    # Heading
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    
    # Content
    para = doc.add_paragraph(content)
    para.runs[0].font.size = Pt(10)
    
    # Add spacing
    doc.add_paragraph()


def generate_resume_with_template(data: dict, filename: str, template: str = "modern"):
    """
    Generate resume with different templates (future enhancement)
    
    Args:
        data: Structured resume data
        filename: Output file path
        template: Template style ("modern", "classic", "minimalist")
    """
    # For now, using the default ATS-friendly template
    # Can be extended with different styles
    generate_ats_resume(data, filename)


def validate_resume_data(data: dict) -> tuple[bool, list]:
    """
    Validate resume data before generation
    
    Args:
        data: Resume data to validate
        
    Returns:
        Tuple of (is_valid, errors)
    """
    errors = []
    
    # Check required top-level keys
    required_keys = ['personal', 'skills', 'experience', 'education']
    for key in required_keys:
        if key not in data:
            errors.append(f"Missing required section: {key}")
    
    # Check personal information
    if 'personal' in data:
        personal = data['personal']
        if not personal.get('name'):
            errors.append("Name is required in personal section")
        if not personal.get('email'):
            errors.append("Email is recommended in personal section")
    
    # Check if at least one content section has data
    has_content = (
        len(data.get('skills', [])) > 0 or
        len(data.get('experience', [])) > 0 or
        len(data.get('education', [])) > 0 or
        data.get('summary', '').strip() != ''
    )
    
    if not has_content:
        errors.append("Resume must have at least one content section with data")
    
    is_valid = len(errors) == 0
    return is_valid, errors


def get_resume_statistics(data: dict) -> dict:
    """
    Get statistics about the resume
    
    Args:
        data: Resume data
        
    Returns:
        Dictionary with statistics
    """
    return {
        'total_skills': len(data.get('skills', [])),
        'total_experiences': len(data.get('experience', [])),
        'total_projects': len(data.get('projects', [])),
        'total_education': len(data.get('education', [])),
        'total_certifications': len(data.get('certifications', [])),
        'total_achievements': len(data.get('achievements', [])),
        'has_summary': bool(data.get('summary', '').strip()),
        'has_contact_info': bool(data.get('personal', {}).get('email')),
    }